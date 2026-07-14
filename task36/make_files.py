"""Generates the four supporting attachments from the shared data model.

House rules enforced here, from the project's poor-quality-artifact guidance:
no "AI blue" palette, no italic subheadings or footers, no confidentiality
notices, no mention of the project, and no input file that names another input
file. Where one document has to refer to another it does so by document class
("the plant record", "a field change notice"), never by filename.

Filename conventions differ between files on purpose: in a real network these
come out of four different systems.
"""
import csv
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side

from model import (SPEC, TIER_WEIGHT, FIELD_CHANGE, LEGS_AS_BUILT,
                   otdr_events, ont_rows)

BLACK = colors.black
GREY = colors.HexColor("#4a4a4a")
BAND = colors.HexColor("#e6e6e6")
RULE = colors.HexColor("#8c8c8c")

# ======================================================= 1. loss budget standard
doc = SimpleDocTemplate("attachments/OSP-207_RevE_Optical_Loss_Budget.pdf",
                        pagesize=LETTER, topMargin=0.8*inch, bottomMargin=0.8*inch,
                        leftMargin=0.9*inch, rightMargin=0.9*inch,
                        title="OSP-207 Rev E — Optical Loss Budget Standard",
                        author="Outside Plant Engineering")
ss = getSampleStyleSheet()
title = ParagraphStyle("t", parent=ss["Normal"], fontName="Helvetica-Bold",
                       fontSize=15, textColor=BLACK, spaceAfter=2, leading=19)
docline = ParagraphStyle("d", parent=ss["Normal"], fontName="Helvetica",
                         fontSize=9.5, textColor=GREY, spaceAfter=14)
h2 = ParagraphStyle("h2", parent=ss["Normal"], fontName="Helvetica-Bold",
                    fontSize=11, textColor=BLACK, spaceBefore=15, spaceAfter=6)
body = ParagraphStyle("b", parent=ss["Normal"], fontName="Helvetica",
                      fontSize=9.8, leading=14, textColor=BLACK)


def tbl(data, widths):
    t = Table(data, colWidths=widths, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BAND),
        ("TEXTCOLOR", (0, 0), (-1, -1), BLACK),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.2),
        ("GRID", (0, 0), (-1, -1), 0.5, RULE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
    ]))
    return t


S = [
    Paragraph("OPTICAL LOSS BUDGET STANDARD", title),
    Paragraph("Standard OSP-207&nbsp;&nbsp;|&nbsp;&nbsp;Revision E&nbsp;&nbsp;|&nbsp;&nbsp;"
              "Effective 20 February 2026&nbsp;&nbsp;|&nbsp;&nbsp;Owner: Outside Plant "
              "Engineering&nbsp;&nbsp;|&nbsp;&nbsp;Supersedes Revision D", docline),

    Paragraph("1. Scope", h2),
    Paragraph("This standard sets the component loss values, the received-power acceptance "
              "thresholds, and the remediation priority rules used to verify Gigabit Passive "
              "Optical Network (GPON) distribution plant. It applies at the 1490 nm downstream "
              "wavelength. All fibre in the access plant is ITU-T G.652.D single mode. Revision E "
              "withdraws the nominal 1:8 splitter allowance and introduces the SLA credit "
              "exposure rule in section 8.", body),

    Paragraph("2. Component loss values", h2),
    Paragraph("These are design values. A component whose measured loss exceeds its design value "
              "is out of specification and shall be recorded as such.", body),
    Spacer(1, 7),
    tbl([
        ["Component", "Design loss", "Unit"],
        ["Single-mode fibre attenuation at 1490 nm", f"{SPEC['fiber_db_per_km']:.2f}", "dB per km"],
        ["Fusion splice", f"{SPEC['fusion_splice_db']:.2f}", "dB each"],
        ["Mechanical splice", f"{SPEC['mechanical_splice_db']:.2f}", "dB each"],
        ["Mated connector pair", f"{SPEC['connector_pair_db']:.2f}", "dB each"],
        ["1:4 optical splitter, insertion loss", f"{SPEC['splitter_1x4_db']:.2f}", "dB"],
        ["1:8 optical splitter, insertion loss", "certified per port", "see section 3"],
    ], [3.5*inch, 1.35*inch, 1.35*inch]),

    Paragraph("3. Primary splitter insertion loss", h2),
    Paragraph("<b>No nominal insertion loss shall be assumed for a 1:8 primary splitter.</b> Each "
              "1:8 splitter is individually certified at manufacture and its ports differ. The "
              "insertion loss that applies to a distribution leg is the certified loss of the "
              "specific hub port that feeds that leg, as recorded in the plant record for the hub. "
              "Using an averaged or nominal figure in place of the certified port value is a "
              "reportable budgeting error.", body),

    Paragraph("4. Transmit reference", h2),
    Paragraph(f"Optical line terminal (OLT) transmit power at the passive optical network (PON) "
              f"port is commissioned and held at <b>{SPEC['olt_tx_dbm']:+.2f} dBm</b>. End-to-end "
              "path loss is the difference between that transmit power and the power received at "
              "the optical network terminal (ONT).", body),

    Paragraph("5. Received-power acceptance thresholds", h2),
    Paragraph("Every in-service ONT is classified from its measured received power as set out "
              "below. Boundaries are inclusive as written.", body),
    Spacer(1, 7),
    tbl([
        ["Classification", "Measured ONT received power"],
        ["PASS", f"greater than or equal to {SPEC['pass_min_dbm']:.2f} dBm"],
        ["MARGINAL", f"below {SPEC['pass_min_dbm']:.2f} dBm and greater than or equal to "
                     f"{SPEC['fail_below_dbm']:.2f} dBm"],
        ["FAIL", f"below {SPEC['fail_below_dbm']:.2f} dBm"],
    ], [1.5*inch, 4.3*inch]),

    Paragraph("6. Measurement tolerance", h2),
    Paragraph(f"Power meters and optical time-domain reflectometers (OTDRs) in field service carry "
              f"a combined uncertainty of plus or minus {SPEC['tolerance_db']:.2f} dB. A difference "
              f"between measured loss and design loss of {SPEC['tolerance_db']:.2f} dB or less is "
              "within tolerance and shall not be raised as a fault. A difference greater than "
              f"{SPEC['tolerance_db']:.2f} dB indicates excess loss physically present in the plant "
              "and shall be investigated.", body),

    Paragraph("7. OTDR test method for distribution legs", h2),
    Paragraph("A distribution leg is tested by shooting an OTDR from the fibre distribution hub "
              "(FDH) output port toward the multiport service terminal (MST). The trace datum, "
              "0.00 km, is the FDH output port. Because the MST houses a splitter, the trace "
              "terminates at the MST input connector: <b>it cannot see any fibre, connector or "
              "splice downstream of the MST, and that includes all subscriber drop cable.</b> "
              "Excess loss downstream of an MST is therefore invisible to this test and must be "
              "isolated from received-power measurements taken at the subscriber terminal. An OTDR "
              "reports an event by its distance from the datum and does not report whether a splice "
              "is fusion or mechanical; the splice type is established from the plant record, and "
              "the design value for that type is the one that applies.", body),

    Paragraph("8. Remediation priority", h2),
    Paragraph("Faults are ranked for remediation by severity first, placing FAIL ahead of MARGINAL. "
              "Faults of equal severity are then ranked by <b>SLA credit exposure</b>, in "
              "descending order. Any remaining tie is broken by the lowest measured received power. "
              "SLA credit exposure is the sum of the service-tier weights of every subscriber "
              "affected by the fault. Subscriber headcount alone is not a priority measure: a "
              "single high-tier subscriber can outrank an entire leg of residential service.", body),
    Spacer(1, 7),
    tbl([["Service tier", "SLA credit weight"]] +
        [[k, str(v)] for k, v in sorted(TIER_WEIGHT.items(), key=lambda kv: kv[1])],
        [3.0*inch, 1.8*inch]),

    Paragraph("9. Change control", h2),
    Paragraph("A field change notice supersedes the as-built plant record for the segment it names, "
              "with effect from its issue date. Where a field change notice is in force, the "
              "geometry, splice schedule and component counts it specifies replace those in the "
              "as-built record for that segment. All other segments continue to be governed by the "
              "as-built record. Loss budgets computed against superseded geometry are invalid.", body),
]
doc.build(S)
print("wrote attachments/OSP-207_RevE_Optical_Loss_Budget.pdf")

# ============================================================ 2. field change notice
d = Document()
for s in d.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)

st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.font.color.rgb = RGBColor(0, 0, 0)

h = d.add_paragraph()
r = h.add_run("FIELD CHANGE NOTICE")
r.bold = True
r.font.size = Pt(16)

sub = d.add_paragraph()
r = sub.add_run(f"Notice {FIELD_CHANGE['ref']}     Issued {FIELD_CHANGE['issued']}     "
                "Outside Plant Engineering     Status: IN FORCE")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

d.add_paragraph()

meta = d.add_table(rows=4, cols=2)
meta.style = "Table Grid"
lg = FIELD_CHANGE["leg"]
for i, (k, v) in enumerate([
    ("Hub", "FDH-14, Ardwick Exchange (ARD-01)"),
    ("Segment affected", f"Distribution Leg {lg}, FDH-14 to {LEGS_AS_BUILT[lg]['mst']}"),
    ("Segments unaffected", "Feeder; Distribution Legs A, B, D and E; all subscriber drops"),
    ("Effect", "Supersedes the as-built record for the affected segment only"),
]):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v
    for p in meta.rows[i].cells[0].paragraphs:
        for run in p.runs:
            run.bold = True

d.add_paragraph()
p = d.add_paragraph()
p.add_run("1. Reason for change").bold = True
d.add_paragraph(FIELD_CHANGE["reason"])

p = d.add_paragraph()
p.add_run("2. Revised geometry").bold = True
d.add_paragraph(
    f"The route length of Distribution Leg {lg}, measured from the FDH-14 output port to the "
    f"{LEGS_AS_BUILT[lg]['mst']} input connector, is now "
    f"{FIELD_CHANGE['new_length_km']:.2f} km. The hub port feeding the leg is unchanged, as is "
    f"the mated connector complement. The splice schedule for the leg is replaced in full by the "
    f"schedule below; the splices recorded against this leg in the as-built record no longer exist.")

d.add_paragraph()
t = d.add_table(rows=1 + len(FIELD_CHANGE["new_splices"]), cols=3)
t.style = "Table Grid"
for j, head in enumerate(["Splice", "Distance from FDH-14 (km)", "Splice type"]):
    c = t.rows[0].cells[j]
    c.text = head
    for p2 in c.paragraphs:
        for run in p2.runs:
            run.bold = True
for i, (dist, kind) in enumerate(FIELD_CHANGE["new_splices"], start=1):
    t.rows[i].cells[0].text = str(i)
    t.rows[i].cells[1].text = f"{dist:.2f}"
    t.rows[i].cells[2].text = kind.capitalize()

d.add_paragraph()
p = d.add_paragraph()
p.add_run("3. Records").bold = True
d.add_paragraph(
    "The as-built drawing has not yet been reissued and continues to show the abandoned route for "
    "this leg. Until it is reissued, this notice is the controlling record for the affected "
    "segment. Any optical loss budget computed for subscribers on this leg must be built on the "
    "geometry and splice schedule set out above.")

d.add_paragraph()
f = d.add_paragraph()
r = f.add_run("Outside Plant Engineering     Ardwick Exchange (ARD-01)")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
f.alignment = WD_ALIGN_PARAGRAPH.LEFT

d.save("attachments/FCN-2291 Leg C Rediversion.docx")
print("wrote attachments/FCN-2291 Leg C Rediversion.docx")

# ============================================================== 3. OTDR export
with open("attachments/otdr_export_fdh14_20260302.csv", "w", newline="") as f:
    w = csv.writer(f)
    w.writerow(["trace_id", "leg", "event_no", "distance_from_fdh_km",
                "event_type", "measured_loss_db", "reflectance_db"])
    for e in otdr_events():
        w.writerow([f"FDH14-{e['leg']}", e["leg"], e["event"], f"{e['distance_km']:.2f}",
                    e["type"], f"{e['loss_db']:.2f}", e["reflectance_db"]])
print("wrote attachments/otdr_export_fdh14_20260302.csv")

# ============================================================ 4. ONT power survey
wb = Workbook()
ws = wb.active
ws.title = "Power Survey"
thin = Side(style="thin", color="8C8C8C")
bd = Border(left=thin, right=thin, top=thin, bottom=thin)

ws["A1"] = "ONT Power Survey — FDH-14"
ws["A1"].font = Font(bold=True, size=13, color="000000")
ws["A2"] = "Downstream 1490 nm. Survey window 24 February to 2 March 2026."
ws["A2"].font = Font(size=10, color="4A4A4A")
ws["A3"] = "Received power measured at the ONT optical port with a calibrated power meter."
ws["A3"].font = Font(size=10, color="4A4A4A")

headers = ["ONT ID", "Leg", "Served From", "Drop Length (m)",
           "Measured Received Power (dBm)", "Service Tier"]
for c, htxt in enumerate(headers, start=1):
    cell = ws.cell(row=5, column=c, value=htxt)
    cell.font = Font(bold=True, size=10, color="000000")
    cell.border = bd
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

for i, r in enumerate(ont_rows(), start=6):
    vals = [r["ont"], r["leg"], r["mst"], r["drop_m"], r["measured_rx"], r["tier"]]
    for c, v in enumerate(vals, start=1):
        cell = ws.cell(row=i, column=c, value=v)
        cell.font = Font(size=10)
        cell.border = bd
        if c in (4, 5):
            cell.alignment = Alignment(horizontal="right")
            if c == 5:
                cell.number_format = "0.00"
        else:
            cell.alignment = Alignment(horizontal="left")

for col, wdt in zip("ABCDEF", [13, 8, 14, 16, 30, 24]):
    ws.column_dimensions[col].width = wdt
ws.row_dimensions[5].height = 32
ws.freeze_panes = "A6"
wb.save("attachments/ONT_Power_Survey_FDH14.xlsx")
print("wrote attachments/ONT_Power_Survey_FDH14.xlsx")

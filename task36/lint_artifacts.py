"""Scans every attachment (and the prompt) for the tells listed in the project's
"AI / Poor Image Quality Examples" guidance. Run before uploading anything."""
import glob
import os
import re
import subprocess
import sys

from openpyxl import load_workbook

ATT = "attachments"

# Filenames of the other inputs — an input file must never name another input file.
INPUT_NAMES = [os.path.basename(p) for p in glob.glob(f"{ATT}/*")]
INPUT_STEMS = [os.path.splitext(n)[0] for n in INPUT_NAMES]

AI_NAMES = [
    "sarah", "emily", "emma", "jessica", "michael", "john smith", "jane doe",
    "john doe", "dr. sarah", "chen", "priya", "aisha", "acme", "techcorp",
    "globex", "initech", "nexgen", "meridian", "stellar solutions", "synergy",
    "quantum", "vertex", "apex", "pinnacle", "horizon", "nexus", "zenith",
    "techflow", "datacorp", "cloudnine", "springfield", "anytown", "riverside",
    "riverbend", "greenfield", "fictionville", "example.com", "555-01",
    "123 main", "elm street", "oak avenue", "maple",
]

BANNED_PHRASES = [
    "project rose", "structured visual artifact", "synthetic", "simulated",
    "for internal use only", "strictly confidential", "confidential",
    "do not distribute", "this document was prepared by", "as an ai",
    "certainly!", "it is important to note", "in today's fast-paced",
    "deliverable", "input file", "attached file",
]

AI_BLUE_HEX = ["1f77b4", "4472c4", "2f6f9f", "1b2733", "5b6b7a", "e9eff5",
               "4e79a7", "0f4c81", "003366", "1f3864"]


def text_of_pdf(path):
    try:
        out = subprocess.run(["python3", "-c",
                              f"from pypdf import PdfReader;"
                              f"print(''.join(p.extract_text() for p in PdfReader({path!r}).pages))"],
                             capture_output=True, text=True, timeout=60)
        return out.stdout
    except Exception:
        return ""


def text_of_xlsx(path):
    wb = load_workbook(path)
    chunks = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if v is not None:
                    chunks.append(str(v))
    return "\n".join(chunks)


def text_of_docx(path):
    from docx import Document
    d = Document(path)
    chunks = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            chunks += [c.text for c in row.cells]
    return "\n".join(chunks)


def text_of(path):
    if path.endswith(".pdf"):
        return text_of_pdf(path)
    if path.endswith(".xlsx"):
        return text_of_xlsx(path)
    if path.endswith(".docx"):
        return text_of_docx(path)
    if path.endswith((".csv", ".txt")):
        return open(path, encoding="utf-8", errors="replace").read()
    return ""


issues = []

for path in sorted(glob.glob(f"{ATT}/*")) + ["prompt.txt"]:
    name = os.path.basename(path)
    body = text_of(path)
    low = body.lower()

    for other_stem, other_name in zip(INPUT_STEMS, INPUT_NAMES):
        if other_name == name:
            continue
        if other_stem.lower() in low or other_name.lower() in low:
            # The prompt is *required* to name the input files; inputs are not.
            if name != "prompt.txt":
                issues.append(f"{name}: names another input file ({other_name})")

    for tell in AI_NAMES:
        if tell in low:
            issues.append(f"{name}: AI-default name/entity '{tell}'")

    for phrase in BANNED_PHRASES:
        if phrase in low and name != "prompt.txt":
            issues.append(f"{name}: banned phrase '{phrase}'")

    if name == "prompt.txt":
        if re.search(r"\*\*|^#{1,6} |^\s*[-*]\s", body, re.M):
            issues.append("prompt.txt: markdown formatting (flagged as an AI tell)")

    # future-dated content
    for yr in re.findall(r"\b20(2[7-9]|[3-9]\d)\b", body):
        issues.append(f"{name}: future-dated content (20{yr})")

# palette check on generator sources
for src in ["make_sva.py", "make_files.py"]:
    s = open(src).read().lower()
    for hexv in AI_BLUE_HEX:
        if hexv in s:
            issues.append(f"{src}: 'AI blue' palette colour #{hexv}")

print(f"attachments: {', '.join(sorted(INPUT_NAMES))}\n")
if issues:
    print(f"{len(issues)} ISSUE(S):")
    for i in sorted(set(issues)):
        print("  -", i)
    sys.exit(1)
print("clean — no AI tells, no cross-file references, no banned palette")
